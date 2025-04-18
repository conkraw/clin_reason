import streamlit as st
import pandas as pd
import os
import glob
import random
import datetime
import re
import openai 

from docx import Document
from docx.shared import Inches

import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders

import firebase_admin
from firebase_admin import credentials, firestore

import streamlit.components.v1 as components

# Set wide layout
st.set_page_config(layout="wide")

openai.api_key = st.secrets["openai"]["api_key"]

# Initialize Firebase
firebase_creds = st.secrets["firebase_service_account"].to_dict()
if not firebase_admin._apps:
    cred = credentials.Certificate(firebase_creds)
    firebase_admin.initialize_app(cred)
db = firestore.client()

# Session state initialization
def initialize_state():
    keys = ["authenticated", "user_name", "assigned_passcode", "recipient_email", 
            "question_row", "selected_diagnoses", "search_input", "answered", "review_sent"]
    for key in keys:
        if key not in st.session_state:
            if key in ["authenticated", "answered", "review_sent"]:
                st.session_state[key] = False
            else:
                st.session_state[key] = ""
    if "search_input_key" not in st.session_state:
        st.session_state.search_input_key = 0
    if "clear_search" not in st.session_state:
        st.session_state.clear_search = False

def lock_passcode_if_needed():
    passcode_str = str(st.session_state.assigned_passcode).strip()
    if not passcode_str:
        st.error("Assigned passcode is empty. Please log in with a valid passcode.")
        st.stop()  # or return False, depending on your flow
    if passcode_str.lower() == "password":
        return False  # Allow default password

    doc_ref = db.collection("shelf_records_prioritized").document(passcode_str)
    now = datetime.datetime.now(datetime.timezone.utc)
    doc = doc_ref.get()
    if doc.exists:
        data = doc.to_dict()
        ts = data.get("timestamp")
        if ts is not None:
            try:
                ts_dt = ts.to_datetime()
            except AttributeError:
                ts_dt = ts.replace(tzinfo=datetime.timezone.utc)
            if (now - ts_dt).total_seconds() < 6 * 3600:
                st.session_state.lock_timestamp = ts_dt
                return True
    doc_ref.set({"processed": True, "timestamp": firestore.SERVER_TIMESTAMP})
    st.session_state.lock_timestamp = now
    return False

    
def check_and_add_passcode(passcode):
    passcode_str = str(passcode).strip()
    if not passcode_str:
        return False  # or alternatively, raise an error if you prefer.
    if passcode_str.lower() == "password":
        return False  # Allow default password
    
    doc_ref = db.collection("shelf_records_prioritized").document(passcode_str)
    doc = doc_ref.get()
    now = datetime.datetime.now(datetime.timezone.utc)
    if doc.exists:
        data = doc.to_dict()
        ts = data.get("timestamp")
        if ts is not None:
            try:
                ts_dt = ts.to_datetime()
            except AttributeError:
                ts_dt = ts.replace(tzinfo=datetime.timezone.utc)
            if (now - ts_dt).total_seconds() < 6 * 3600:
                return True
    doc_ref.set({"processed": True, "timestamp": firestore.SERVER_TIMESTAMP})
    return False


def send_email_with_attachment(to_emails, subject, body, attachment_path):
    from_email = st.secrets["general"]["email"]
    password = st.secrets["general"]["email_password"]
    msg = MIMEMultipart()
    msg['From'] = from_email
    msg['To'] = ', '.join(to_emails)
    msg['Subject'] = subject
    msg.attach(MIMEText(body, 'html'))
    with open(attachment_path, 'rb') as attachment:
        part = MIMEBase('application', 'octet-stream')
        part.set_payload(attachment.read())
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', 'attachment', filename=os.path.basename(attachment_path))
        msg.attach(part)
    try:
        with smtplib.SMTP_SSL('smtp.gmail.com', 465) as server:
            server.login(from_email, password)
            server.send_message(msg, from_addr=from_email, to_addrs=to_emails)
        st.success("Email sent successfully!")
    except Exception as e:
        st.error(f"Error sending email: {e}")

def format_physical_exam(pe_text):
    if not isinstance(pe_text, str):
        return []
    pattern = r'([A-Z][a-zA-Z ]+):'
    parts = re.split(pattern, pe_text)
    formatted_lines = []
    for i in range(1, len(parts), 2):
        label = parts[i].strip()
        description = parts[i + 1].strip()
        formatted_lines.append(f"{label}: {description}")
    return formatted_lines

def safe_text(val):
    return str(val) if pd.notna(val) else ""

def display_pretty_table(user_order, correct_order):
    html_table = """
    <style>
    table.custom-table {
      width: 70%;
      margin: 1rem auto;
      border-collapse: collapse;
      font-family: Arial, sans-serif;
      font-size: 15px;
    }
    table.custom-table th {
      background-color: #f0f0f0;
      text-align: left;
      padding: 8px;
      border: 1px solid #ccc;
    }
    table.custom-table td {
      border: 1px solid #ccc;
      padding: 8px;
    }
    </style>
    <table class="custom-table">
      <thead>
        <tr>
          <th>Rank</th>
          <th>Student Answers</th>
          <th>Correct Answers</th>
        </tr>
      </thead>
      <tbody>
    """
    for i, (ua, ca) in enumerate(zip(user_order, correct_order), start=1):
        html_table += f"""
          <tr>
            <td>{i}</td>
            <td>{ua}</td>
            <td>{ca}</td>
          </tr>
        """
    html_table += """
      </tbody>
    </table>
    """
    components.html(html_table, height=250)

def get_used_cases_for_preceptor(designation):
    """Fetches record_ids used in the last 7 days for a given preceptor designation."""
    used_cases = []
    collection_name = "global_used_cases_" + designation if designation else "global_used_cases"
    used_ref = db.collection(collection_name)
    docs = used_ref.stream()
    now = datetime.datetime.utcnow()
    for doc in docs:
        data = doc.to_dict()
        ts = data.get("timestamp")
        if ts is not None:
            ts_naive = ts.replace(tzinfo=None)
            if (now - ts_naive).days < 7:
                used_cases.append(doc.id)
            else:
                doc.reference.delete()
    return used_cases

def mark_case_as_used_for_preceptor(designation, record_id):
    """Marks a given record_id as used for the specified preceptor designation."""
    collection_name = "global_used_cases_" + designation if designation else "global_used_cases"
    used_ref = db.collection(collection_name)
    used_ref.document(str(record_id)).set({
         "used": True,
         "timestamp": firestore.SERVER_TIMESTAMP
    })

def generate_review_doc_prioritized(row, user_order, output_filename="review.docx"):
    doc = Document()
    doc.add_heading("Review of Incorrect Prioritized Diagnosis", level=1)
    doc.add_heading(f"Student: {st.session_state.user_name}", level=2)
    doc.add_heading(f"Case ({row['record_id']}):", level=2)
    doc.add_paragraph(safe_text(row["anchorx"]))
    sections = {
         "Chief Complaint": row.get("cc", ""),
         "History of Present Illness": row.get("hpi", ""),
         "Past Medical History": row.get("pmhx", ""),
         "Medications": row.get("meds", ""),
         "Allergies": row.get("allergies", ""),
         "Immunizations": row.get("immunizations", ""),
         "Social History": row.get("shx", ""),
         "Family History": row.get("fhx", ""),
         "Vital Signs": row.get("vs", ""),
         "Physical Exam": row.get("pe", "")
    }
    for title, content in sections.items():
         if pd.notna(content) and str(content).strip():
              doc.add_heading(title, level=2)
              if title == "Physical Exam":
                  lines = format_physical_exam(content)
                  for line in lines:
                      if ":" in line:
                          label, text = line.split(":", 1)
                          p = doc.add_paragraph()
                          run1 = p.add_run(f"{label.strip()}: ")
                          run1.bold = True
                          p.add_run(text.strip())
                      else:
                          doc.add_paragraph(line)
              else:
                  doc.add_paragraph(safe_text(content))
    doc.add_heading("Student Prioritized Diagnosis:", level=2)
    for i, diag in enumerate(user_order):
         doc.add_paragraph(f"{i+1}. {diag}")
    correct_order = [safe_text(row.get("answer", "")).strip(), 
                     safe_text(row.get("sec_dx", "")).strip(), 
                     safe_text(row.get("thir_dx", "")).strip()]
    doc.add_heading("Correct Prioritized Diagnosis:", level=2)
    for i, diag in enumerate(correct_order):
         doc.add_paragraph(f"{i+1}. {diag}")
    doc.add_heading("Explanation:", level=2)
    doc.add_paragraph(safe_text(row.get("answer_explanationx", "")))
    doc.save(output_filename)
    return output_filename

def get_best_matching_diagnosis(user_input, choices, case_anchor=""):
    """
    Uses OpenAI's ChatCompletion API to select the diagnosis from choices that is semantically
    closest to the user input—even if the input is a partial word.
    
    The function uses the provided case context (case_anchor) to guide the decision.
    If none of the diagnoses appears appropriate, it will return "No suitable match".
    
    Returns:
        A string with the diagnosis exactly as it appears in the choices, or "No suitable match".
    """
    # Create a prompt with explicit instructions:
    prompt = (
        f"You are an expert medical assistant. Here is a list of possible diagnoses: {', '.join(choices)}. "
        f"The clinical scenario is described as: \"{case_anchor}\". "
        f"A user has typed in the query: \"{user_input}\". "
        "Even if the input is only a fragment (for example, a partial word), "
        "please select the diagnosis from the provided list that is the best semantic match to the input. "
        "Return only the diagnosis exactly as it appears in the list. If none of the provided diagnoses fit, "
        "reply with exactly 'No suitable match'."
    )

    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.0,  # Lower temperature for a deterministic response.
        )
        answer = response["choices"][0]["message"]["content"].strip()
        return answer
    except Exception as e:
        st.error(f"Error obtaining AI suggestion: {e}")
        return None


def get_clinical_context(row):
    context_parts = []
    for field in ["cc", "hpi", "pmhx", "meds", "allergies", "immunizations", "shx", "fhx", "vs", "pe"]:
        value = row.get(field, "")
        # Convert the value to a string first so that None becomes "None"
        # Alternatively, check explicitly for None.
        if value is not None:
            value_str = str(value).strip()
        else:
            value_str = ""
        if value_str:
            context_parts.append(f"{field.upper()}: {value_str}")
    return " | ".join(context_parts)

def save_prioritized_exam_state():
    # Save only the fields that are needed to resume the exam.
    data = {
        "question_row": st.session_state.question_row,
        "selected_diagnoses": st.session_state.selected_diagnoses,
        "answered": st.session_state.answered,
        "review_sent": st.session_state.review_sent,
        "last_search_query": st.session_state.get("last_search_query", ""),
        "ai_suggestion": st.session_state.get("ai_suggestion", ""),
        # Add additional keys as needed.
        "timestamp": firestore.SERVER_TIMESTAMP,
    }
    user_key = str(st.session_state.assigned_passcode)
    db.collection("exam_sessions_prioritized").document(user_key).set(data)

def load_prioritized_exam_state():
    user_key = str(st.session_state.assigned_passcode)
    doc_ref = db.collection("exam_sessions_prioritized").document(user_key)
    doc = doc_ref.get()
    if doc.exists:
        data = doc.to_dict()
        st.session_state.question_row = data.get("question_row", "")
        st.session_state.selected_diagnoses = data.get("selected_diagnoses", [])
        st.session_state.answered = data.get("answered", False)
        st.session_state.review_sent = data.get("review_sent", False)
        st.session_state.last_search_query = data.get("last_search_query", "")
        st.session_state.ai_suggestion = data.get("ai_suggestion", "")

def lock_passcode_on_submission(passcode):
    passcode_str = str(passcode).strip()
    if not passcode_str:
        return
    doc_ref = db.collection("shelf_records_prioritized").document(passcode_str)
    doc_ref.set({
        "processed": True, 
        "timestamp": firestore.SERVER_TIMESTAMP, 
        "locked": True
    })

def is_passcode_locked(passcode):
    """
    Checks whether the given passcode is locked.
    A passcode is considered locked if its document has "locked": True and the stored timestamp
    is less than 6 hours old.
    """
    passcode_str = str(passcode).strip()
    if not passcode_str:
        return False
    doc_ref = db.collection("shelf_records_prioritized").document(passcode_str)
    doc = doc_ref.get()
    now = datetime.datetime.now(datetime.timezone.utc)
    if doc.exists:
        data = doc.to_dict()
        if data.get("locked", False):
            ts = data.get("timestamp")
            if ts is not None:
                try:
                    ts_dt = ts.to_datetime()
                except AttributeError:
                    ts_dt = ts.replace(tzinfo=datetime.timezone.utc)
                if (now - ts_dt).total_seconds() < 6 * 3600:
                    return True
    return False

def save_completed_exam():
    """
    Saves the completed exam details permanently in Firestore.
    Stores the passcode used, record_id of the question, student name, answers provided,
    and a timestamp.
    """
    user_key = str(st.session_state.assigned_passcode)
    completed_data = {
        "passcode": user_key,
        "student_name": st.session_state.user_name,
        "record_id": st.session_state.question_row.get("record_id", ""),
        "selected_diagnoses": st.session_state.selected_diagnoses,
        "timestamp": firestore.SERVER_TIMESTAMP,
    }
    db.collection("completed_exam_sessions").document().set(completed_data)
    
# Login Screen
def login_screen():
    st.title("Shelf Examination Login")
    passcode_input = st.text_input("Enter your assigned passcode", type="password")
    user_name = st.text_input("Enter your name")
    
    if st.button("Login"):
        if not passcode_input.strip():
            st.error("Please enter your assigned passcode.")
            return
        if "recipients" not in st.secrets:
            st.error("Recipient emails not configured.")
            return
        if passcode_input not in st.secrets["recipients"]:
            st.error("Invalid passcode.")
            return
        if not user_name.strip():
            st.error("Please enter your name.")
            return

        if is_passcode_locked(passcode_input.strip()):
            st.error("This passcode has been used recently. Please try again after 6 hours.")
            st.stop() 
            
        st.session_state.assigned_passcode = passcode_input.strip()
        st.session_state.user_name = user_name.strip()
        st.session_state.recipient_email = st.secrets["recipients"][passcode_input]
        st.session_state.authenticated = True

        load_prioritized_exam_state()
        
        st.rerun()


        
# Prioritized Differential Diagnosis Exam Screen
def exam_screen_prioritized():
    st.title("Shelf Examination – Prioritized Differential Diagnosis")
    st.write(f"Welcome, {st.session_state.user_name}!")

    #if not st.session_state.get("exam_initialized", False):
    #    if lock_passcode_if_needed():
    #        st.error("This passcode has been used recently. Please try again after 6 hours.")
    #        st.stop()  # Stop further processing.
    #    else:
    #        st.session_state.exam_initialized = True


    # 1) LOAD A RANDOM CASE IF NOT ALREADY LOADED
    if not st.session_state.question_row:
        csv_files = glob.glob("*.csv")
        #csv_files = [file for file in csv_files if "prioritized" in file.lower()]
        df_list = [pd.read_csv(file) for file in csv_files]
        df = pd.concat(df_list, ignore_index=True)

        # Extract designation from password (e.g., password1_aaa yields "aaa")
        password = st.session_state.assigned_passcode
        designation = password.split("_")[-1] if "_" in password else ""

        used_cases = get_used_cases_for_preceptor(designation)
        available_df = df[~df["record_id"].isin(used_cases)]

        if available_df.empty:
            st.error("No further cases available for your preceptor at this time. Please try again later.")
            st.stop()
            
        # Sample one case from the available cases
        selected = available_df.sample(1).iloc[0].to_dict()
        st.session_state.question_row = selected
        st.session_state.selected_diagnoses = []
        st.session_state.search_input = ""
        st.session_state.answered = False
        st.session_state.review_sent = False

        mark_case_as_used_for_preceptor(designation, selected["record_id"])
        
    row = st.session_state.question_row

    # 2) SIDEBAR: Display clinical information in collapsible sections
    with st.sidebar:
        st.header("Clinical Information")
        label_map = {
            "cc": "Chief Complaint",
            "hpi": "History of Present Illness",
            "pmhx": "Past Medical History",
            "meds": "Medications",
            "allergies": "Allergies",
            "immunizations": "Immunizations",
            "shx": "Social History",
            "fhx": "Family History",
            "vs": "Vital Signs",
            "pe": "Physical Exam",
        }
        for key, display_label in label_map.items():
            content = row.get(key, "")
            if pd.notna(content) and str(content).strip():
                with st.expander(display_label, expanded=False):
                    if key == "pe":
                        lines = format_physical_exam(content)
                        for line in lines:
                            st.markdown(f"- {line}")
                    else:
                        st.write(content)

    # 3) MAIN PROMPT
    st.subheader(row.get("anchorx", "Please select and prioritize 3 diagnoses:"))
    st.write("Type to search for a diagnosis, then click to add it to your prioritized list. You can reorder or remove items as needed.")
    
    # 4) DIAGNOSIS SEARCH INPUT
    if st.session_state.get("clear_search", False):
        search_input = st.text_input("Type diagnosis:", value="", key="diag_search_input")
        st.session_state.clear_search = False
    else:
        search_input = st.text_input("Type diagnosis:", key="diag_search_input")
    
    # Parse the list of possible diagnoses.
    all_choices = [c.strip() for c in str(row.get("choices", "")).split(",")]
    
    # Only proceed if the user has typed at least 2 characters.
    if len(search_input) >= 2:
        # Try simple substring matching first.
        matches = [c for c in all_choices if search_input.lower() in c.lower()]
    else:
        matches = []
    
    if matches:
        st.write("Matching diagnoses:")
        for match in matches:
            if match not in st.session_state.selected_diagnoses:
                if st.button(f"➕ {match}", key=f"match_{match}"):
                    st.session_state.selected_diagnoses.append(match)
                    st.session_state.clear_search = True
                    save_prioritized_exam_state() 
                    st.rerun()
    # Only try AI suggestion if search_input is not empty and no matches were found.
    elif search_input:  
        last_query = st.session_state.get("last_search_query", "")
        if last_query != search_input:
            case_context = get_clinical_context(row)  # Use your function to build context, if desired.
            ai_suggestion = get_best_matching_diagnosis(search_input, all_choices, case_anchor=case_context)
            st.session_state["ai_suggestion"] = ai_suggestion
            st.session_state["last_search_query"] = search_input
        else:
            ai_suggestion = st.session_state.get("ai_suggestion", None)
        
        if ai_suggestion and ai_suggestion != "No suitable match":
            st.write("Possible Diagnosis: " + ai_suggestion)
            if st.button(f"➕ {ai_suggestion}", key="ai_suggestion_btn"):
                if ai_suggestion not in st.session_state.selected_diagnoses:
                    st.session_state.selected_diagnoses.append(ai_suggestion)
                    st.session_state.clear_search = True
                    save_prioritized_exam_state()
                    st.rerun()
        else:
            st.write("No suggestion available for the entered input.")
            
    # 5) DISPLAY SELECTED DIAGNOSES WITH REORDER/REMOVE OPTIONS
    st.write("Prioritized Differential Diagnosis:")
    arrow_up = "⬆️"
    arrow_down = "⬇️"
    trash_icon = "🗑️"
    for i, diag in enumerate(st.session_state.selected_diagnoses):
        col1, col2, col3, col4 = st.columns([6, 1, 1, 1])
        with col1:
            st.write(f"{i+1}. {diag}")
        with col2:
            if i > 0:
                if st.button(arrow_up, key=f"up_{i}"):
                    st.session_state.selected_diagnoses[i], st.session_state.selected_diagnoses[i-1] = \
                        st.session_state.selected_diagnoses[i-1], st.session_state.selected_diagnoses[i]
                    save_prioritized_exam_state()
                    st.rerun()
        with col3:
            if i < len(st.session_state.selected_diagnoses) - 1:
                if st.button(arrow_down, key=f"down_{i}"):
                    st.session_state.selected_diagnoses[i], st.session_state.selected_diagnoses[i+1] = \
                        st.session_state.selected_diagnoses[i+1], st.session_state.selected_diagnoses[i]
                    save_prioritized_exam_state()
                    st.rerun()
        with col4:
            if st.button(trash_icon, key=f"remove_{i}"):
                st.session_state.selected_diagnoses.pop(i)
                save_prioritized_exam_state()
                st.rerun()

    # 6) SUBMISSION: Only if exactly 3 diagnoses are selected
    if len(st.session_state.selected_diagnoses) == 3 and not st.session_state.answered:
        if st.button("Submit Answer"):
            st.session_state.answered = True
            correct_order = [
                safe_text(row.get("answer", "")).strip(),
                safe_text(row.get("sec_dx", "")).strip(),
                safe_text(row.get("thir_dx", "")).strip(),
            ]
            user_order = [diag.strip() for diag in st.session_state.selected_diagnoses]
            st.write("**Your Prioritized Diagnosis:**")
            display_pretty_table(user_order, correct_order)

            lock_passcode_on_submission(st.session_state.assigned_passcode)
            
            if user_order == correct_order:
                st.success("Correct!")
            else:
                st.error("Incorrect.")
                st.info(row.get("answer_explanationx", ""))
                if not st.session_state.review_sent:
                    filename = f"review_{st.session_state.user_name}_{row['record_id']}.docx"
                    generate_review_doc_prioritized(row, user_order, filename)
                    send_email_with_attachment(
                        to_emails=[st.session_state.recipient_email],
                        subject="Review of Incorrect Prioritized Diagnosis Answer",
                        body="Please find attached a review of your response.",
                        attachment_path=filename
                    )
                    st.session_state.review_sent = True
            st.success("Case complete. Thank you for your response. You may now close the window.")

            save_completed_exam()
            
            user_key = str(st.session_state.assigned_passcode)
            db.collection("exam_sessions_prioritized").document(user_key).delete()
            
    elif len(st.session_state.selected_diagnoses) != 3 and not st.session_state.answered:
        st.info(f"Please select exactly 3 diagnoses. You have selected {len(st.session_state.selected_diagnoses)}.")

# Main App Logic
def main():
    initialize_state()
    if not st.session_state.authenticated:
        login_screen()
    else:
        exam_screen_prioritized()

if __name__ == "__main__":
    main()

